from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Word document
doc = Document()

# Set the style for the document
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title Page
doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
title = doc.add_heading('Rapport de Projet de Fin d’Études', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Développement d’une Application Statistique avec React, Redux, et Python').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Auteur: [Votre Nom]').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Encadrant: [Nom de l’Encadrant]').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Année Universitaire: 2023/2024').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Université: [Nom de l’Université]').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add page break
doc.add_page_break()

# Abstract
doc.add_heading('Résumé', level=2)
abstract = doc.add_paragraph(
    "Ce rapport présente le développement d'une application statistique utilisant React, Redux, et Python pour le back-end. "
    "L'application permet aux utilisateurs de créer un compte, de se connecter et d'accéder à une interface fournissant des statistiques "
    "sur un ensemble de données contenant des informations sur les ventes et le commerce. "
    "L'application inclut également un modèle d'IA qui propose les meilleures stratégies pour la croissance des affaires."
)
abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Table of Contents
doc.add_heading('Table des Matières', level=2)
table_of_contents = doc.add_paragraph()
run = table_of_contents.add_run("Chapitre 1: Présentation de l’organisme d’accueil et étude de l’existant......")
run.add_tab()
run.add_run("3")
table_of_contents.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph('Chapitre 2: État d’art sur les technologies et solutions existantes....................5')
doc.add_paragraph('Chapitre 3: Présentation de la solution développée.................................................7')
doc.add_paragraph('Chapitre 4: Test et évaluation de la solution................................................................9')

# Add page break
doc.add_page_break()

# Chapter 1
doc.add_heading('Chapitre 1: Présentation de l’organisme d’accueil et étude de l’existant', level=2)

doc.add_heading('1.1 Présentation de l’organisme d’accueil', level=3)
doc.add_paragraph(
    "L’organisme d’accueil pour ce projet est [Nom de l'Entreprise], une entreprise spécialisée dans [Domaine de l'entreprise]. "
    "Elle offre divers services, notamment [Services offerts par l'entreprise]. L'entreprise se distingue par son chiffre d'affaires "
    "de [Chiffre d'affaires] et emploie actuellement [Nombre d'employés] personnes."
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('1.2 Étude de l’existant', level=3)
doc.add_paragraph(
    "Actuellement, l'entreprise utilise [Description de la solution existante ou absence de solution]. "
    "Cette solution présente plusieurs limites, telles que [Critiques et limites]. En l'absence d'une solution optimale, "
    "les conséquences sont [Conséquences], ce qui justifie la nécessité de mettre en place une nouvelle solution."
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Add page break
doc.add_page_break()

# Chapter 2
doc.add_heading('Chapitre 2: État d’art sur les technologies et solutions existantes', level=2)

doc.add_heading('2.1 Présentation des technologies', level=3)
doc.add_paragraph(
    "Dans ce projet, nous avons utilisé React, une bibliothèque JavaScript pour construire des interfaces utilisateur, "
    "ainsi que Redux pour gérer l'état de l'application de manière prévisible. Python a été choisi pour le back-end en raison de sa "
    "flexibilité et de ses puissantes bibliothèques pour le traitement des données et l'IA."
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('2.2 Solutions du marché', level=3)
doc.add_paragraph(
    "Plusieurs autres solutions sont disponibles sur le marché pour le développement d'applications statistiques. "
    "Parmi elles, on trouve [Autres solutions]. Ces solutions présentent des caractéristiques variées, telles que [Description des caractéristiques]. "
    "Un tableau comparatif est présenté ci-dessous pour évaluer ces solutions par rapport à la technologie choisie."
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Add page break
doc.add_page_break()

# Chapter 3
doc.add_heading('Chapitre 3: Présentation de la solution développée', level=2)

doc.add_heading('3.1 Conception et spécification des besoins', level=3)
doc.add_paragraph(
    "La solution proposée est conçue pour répondre aux besoins spécifiques de l'entreprise. "
    "L'architecture du système est divisée en trois couches principales: la couche de présentation, la couche de traitement, "
    "et la couche de données. Chaque couche joue un rôle essentiel pour assurer la modularité et l'efficacité du système."
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('3.2 Développement et configuration', level=3)
doc.add_paragraph(
    "Le développement de l'application a été réalisé en utilisant React pour le front-end, Redux pour la gestion de l'état, et Python pour le back-end. "
    "Le système a été configuré pour intégrer l'authentification des utilisateurs, la visualisation des données statistiques, et un modèle d'IA pour "
    "proposer des stratégies d'amélioration des performances commerciales."
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('3.3 Mise en place', level=3)
doc.add_paragraph(
    "La solution a été déployée dans l'environnement cible avec succès. "
    "L'intégration de la solution a été réalisée en tenant compte des besoins spécifiques de l'entreprise et en optimisant les performances du système."
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Add page break
doc.add_page_break()

# Chapter 4
doc.add_heading('Chapitre 4: Test et évaluation de la solution', level=2)

doc.add_heading('4.1 Tests et validation', level=3)
doc.add_paragraph(
    "Des tests rigoureux ont été effectués pour valider le bon fonctionnement de l'application. "
    "Les tests ont couvert tous les aspects de l'application, y compris la fonctionnalité de recherche, la visualisation des données, "
    "et l'intégration du modèle d'IA."
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('4.2 Évaluation comparative', level=3)
doc.add_paragraph(
    "Une évaluation comparative a été réalisée pour démontrer la valeur ajoutée de la nouvelle solution par rapport à l'existant. "
    "Les résultats montrent une amélioration significative en termes de performance, d'efficacité, et de prise de décision grâce à l'intégration de l'IA."
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Add page break
doc.add_page_break()

# Conclusion
doc.add_heading('Conclusion Générale', level=2)
doc.add_paragraph(
    "En conclusion, le projet a permis de développer une application statistique robuste et efficace répondant aux besoins de l'entreprise. "
    "L'intégration du modèle d'IA apporte une valeur ajoutée significative en optimisant les stratégies commerciales. "
    "Des améliorations futures pourraient inclure l'extension des fonctionnalités de l'application et l'intégration de nouvelles technologies."
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# References
doc.add_heading('Bibliographie', level=2)
doc.add_paragraph(
    "1. [Référence 1]\n"
    "2. [Référence 2]\n"
    "3. [Référence 3]"
).alignment = WD_ALIGN_PARAGRAPH.LEFT

# Save the document
doc.save("/mnt/data/PFE_Report.docx")

